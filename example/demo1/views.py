import os
from dateutil import rrule
from datetime import datetime
import datetime
import time

from django.http import FileResponse
from django.shortcuts import render
from django.shortcuts import redirect
from django.conf import settings
from demo1 import models, email as E
import shutil
from django.db.models import Q
from django.shortcuts import get_object_or_404
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import datetime
CS_key = '666666'


def home(request):
	login_name = ""
	if request.session.get('id'):
		if request.session.get('kind') == 'user':
			user = models.User.objects.get(id=request.session.get('id'))
			login_name = user.username
		else:
			CS = models.Customerserver.objects.get(id=request.session.get('id'))
			login_name = CS.username
	return render(request, 'home.html', {"username": login_name})


def login(request):
	if request.session.get('id') != None:
		return redirect('/')
	if request.method == 'POST':
		username = request.POST.get('username')
		password = request.POST.get('password')
		if request.POST.get('is_CS'):
			if username and password:
				username = username.strip()
				password = password.strip()
				try:
					CS = models.Customerserver.objects.get(username = username)
				except:
					message = '该用户名不存在'
					return render(request, 'login.html', {"message":message})
				if CS.password == password:
					request.session['id'] = CS.id
					request.session['kind'] = "CS"
					message = '登陆成功，欢迎您，' + str(CS.username)
					return render(request, 'main.html', {"message": message, "username": username, "kind": "CS"})			#客服登录成功
				else:
					message = '密码错误'
					return render(request, 'login.html', {"message":message})
			else:
				message = '请输入用户名和密码'
				return render(request, 'login.html', {"message":message})
		elif request.POST.get('is_tech'):
			if username and password:
				username = username.strip()
				password = password.strip()
				try:
					technician = models.Technician.objects.get(username=username)
				except:
					message = '该用户名不存在'
					return render(request, 'login.html', {"message":message})
				if technician.password == password:
					request.session['id'] = technician.id
					request.session['kind'] = "tech"
					message = '登陆成功，欢迎您，' + str(technician.username)
					return render(request, 'main.html', {"message": message, "username": username, "kind": "tech"})			#工单师傅登录成功
				else:
					message = '密码错误'
					return render(request, 'login.html', {"message":message})
			else:
				message = '请输入用户名和密码'
				return render(request, 'login.html', {"message":message})
		else:
			if username and password:
				username = username.strip()
				password = password.strip()
				try:
					user = models.User.objects.get(username = username)
				except:
					message = '该用户名不存在'
					return render(request, 'login.html', {"message":message})
				if user.status == 0:
					message = '该用户尚未验证邮箱'
					return render(request, 'login.html', {"message": message})
				if user.password == password:
					request.session['id'] = user.id
					request.session['kind'] = "user"
					message = '登陆成功，欢迎您，' + str(user.username)
					return render(request, 'main.html', {"message": message, "username": username, "kind": "user"})			#用户登录成功
				else:
					message = '密码错误'
					return render(request, 'login.html', {"message":message})
			else:
				message = '请输入用户名和密码'
				return render(request, 'login.html', {"message":message})
	return render(request, 'login.html')


def register(request):
	if request.session.get('id') != None:  # 只有不在登录状态时才可以进行注册
		return redirect('/')
	if request.method == 'POST':
		username = request.POST.get('username').strip()
		password1 = request.POST.get('password1').strip()
		password2 = request.POST.get('password2').strip()
		real_name = request.POST.get('real_name')
		ID_card = request.POST.get('ID_card')
		email = request.POST.get('email').strip()
		is_CS = request.POST.getlist('is_CS')
		if is_CS:  # 处理客服注册
			key = request.POST.get('CS_key')
			if key != CS_key:
				message = '密钥不正确'
				return render(request, 'register.html', {"message": message})
			else:
				if models.Customerserver.objects.filter(username=username).exists():
					message = '该用户名已被注册'
					return render(request, 'register.html', {"message": message})
				if password1 != password2:
					message = '两次密码不一致'
					return render(request, 'register.html', {"message": message})
				CS = models.Customerserver()
				CS.username = username
				CS.password = password1
				CS.save()
				message = '客服注册成功'
				return render(request, 'home.html', {"message": message})
		legal, message = checkout(username, password1, password2, email)
		if not legal:
			return render(request, 'register.html', {"message": message})
		add_user(username, password1, email, real_name, ID_card)
		message = '注册成功，请查看邮箱验证'
		return render(request, 'home.html', {"message": message})
	return render(request, 'register.html')


def logout(request):
	request.session.flush()
	return redirect('/')


def checkout(username, password1, password2, email):
	legal = True
	message = ''
	if models.User.objects.filter(username=username).exists():
		message = '该用户名已被注册'
		legal = False
	elif models.User.objects.filter(email=email).exists():
		exist = False
		for i in models.User.objects.filter(email=email).all():
			if i.status == 1:
				exist = True
		if exist:
			message = '该邮箱已被注册'
			legal = False
		else:
			legal = True
	if password1 != password2 and legal:
		message = '两次密码不一致'
		legal = False
	return legal, message


def active(request, active_code):
	user =	models.User.objects.filter(active_code=active_code).first()
	if user:
		user.status = 1
		user.save()
		message = '验证成功'
		return render(request, 'home.html', {"message": message})
	else:
		message = '邮箱验证失败，请重新注册'
		return render(request, 'register.html', {"message": message})


def main(request):
	login_name = ""
	kind = ""
	unread = 0
	if request.session.get('id'):
		kind = request.session.get('kind')
		if kind == 'user':
			user = models.User.objects.get(id=request.session.get('id'))
			login_name = user.username
			my_message = models.MyMessage.objects.filter(receiver=user)
			for i in my_message:
				if i.status:
					unread += 1
			return render(request, 'main.html', {"username": login_name, "kind": kind, "unread": unread})
		else:
			CS = models.Customerserver.objects.get(id=request.session.get('id'))
			login_name = CS.username
	return render(request, 'main.html', {"username": login_name, "kind": kind, "unread": unread})


def show_info(request):
	kind = request.session.get('kind')
	if request.session.get('id'):
		if request.session.get('kind') == 'user':
			user = models.User.objects.get(id = request.session.get('id'))
			email = user.email
			real_name = user.real_name
			ID_card = user.ID_card
		else:
			user = models.Customerserver.objects.get(id = request.session.get('id'))
			email = "客服无邮箱记录"
		username = user.username
		password = user.password
	if request.method == 'POST':
		if request.session.get('kind') == 'user':
			email = models.User.objects.get(id = request.session.get('id')).email
			real_name = models.User.objects.get(id=request.session.get('id')).real_name
			ID_card = models.User.objects.get(id = request.session.get('id')).ID_card
			models.User.objects.filter(username=username).delete()
			new_username = request.POST.get('username').strip()
			new_password1 = request.POST.get('password1').strip()
			new_password2 = request.POST.get('password2').strip()
			legal, message = checkout(new_username, new_password1, new_password2, email)
			if not legal:
				user = add_user(username, password, email, real_name, ID_card, False)
				request.session['id'] = user.id
				return render(request, 'show_info.html', locals())
			user = add_user(new_username, new_password1, email, real_name, ID_card, False)
			request.session['id'] = user.id
			username = user.username
			password = user.password
			message = '修改成功'
			kind = request.session['kind']
			return render(request, 'main.html', locals())
		else:
			message = '客服信息一旦创建不可修改'
			username = models.Customerserver.objects.get(id = request.session.get('id')).username
			kind = 'CS'
			return render(request, 'main.html', locals())
	else:
		return render(request, 'show_info.html', locals())


def add_user(username, password, email, real_name, ID_card, new_user=True):
	user = models.User()
	user.username = username
	user.password = password
	user.email = email
	user.status = 1
	user.real_name = real_name
	user.ID_card = ID_card
	if new_user:
		user.status = 0
		user.active_code = E.send_register_email(email)
	user.save()
	return user


def add_house(request):
	if request.method == 'POST':
		house_name = request.POST.get('house_name')
		short_leasing = request.POST.get('short_leasing')
		if short_leasing:
			short_leasing_fee = request.POST.get('short_leasing_fee')
		long_leasing = request.POST.get('long_leasing')
		if long_leasing:
			long_leasing_fee = request.POST.get('long_leasing_fee')
		house_type = request.POST.get('house_type')
		district = request.POST.get('district')
		address = request.POST.get('address')
		contact_number = request.POST.get('contact_number')
		files = request.FILES.getlist("picture", None)
		if files:
			for file in files:
				if not str(file.name).__contains__("jpg"):
					return render(request, 'add_house.html', {"message": "图片仅支持jpg格式"})
		else:
			return render(request, 'add_house.html', {"message": "请上传图片"})

		house = models.House()
		if house_name and house_type and district and address and contact_number and file:
			house.house_name = house_name
			house.short_leasing = short_leasing
			if short_leasing:
				house.short_leasing_fee = short_leasing_fee
			house.long_leasing = long_leasing
			if long_leasing:
				house.long_leasing_fee = long_leasing_fee
			house.house_type = house_type
			house.district = district
			house.address = address
			house.contact_number = contact_number
			house.save()
			house.picture = "/demo1/" + str(house.house_id) + "/"
			os.makedirs(settings.MEDIA_ROOT + "/demo1/" + str(house.house_id) + "/")
			i = 1
			for file in files:
				filename = settings.MEDIA_ROOT + "/demo1/" + str(house.house_id) + "/" + str(i) + ".jpg"
				with open(filename, 'wb') as pic:
					for c in file.chunks():
						pic.write(c)
				i += 1
			house.pic_num = i - 1
			house.save()
			message = "添加成功"
			return render(request, 'main.html', {"message": message, "kind": request.session.get('kind')})
		else:
			message = "请将信息填写完整"
			return render(request, 'add_house.html', {"message": message})
	return render(request, 'add_house.html')


def search_house(request):
	houses = models.House.objects.all()
	houses = houses.exclude(long_leasing=True, status=1)
	login_name = ""
	kind = ""
	if request.session.get('kind') == 'user':
		if request.session.get('id'):
			user = models.User.objects.get(id=request.session.get('id'))
			login_name = user.username
			kind = 'User'
	elif request.session.get('kind') == 'CS':
		if request.session.get('id'):
			cs = models.Customerserver.objects.get(id=request.session.get('id'))
			login_name = cs.username
			kind = 'CS'
	else:
		kind = 'guest'
		login_name = 'guest'
	if request.method == 'POST':
		key_word = request.POST.get('key_word')
		houses = houses.filter(house_name__contains=key_word)
		district = request.POST.get('district')
		if district != '不限':
			houses = houses.filter(district=district)
		short_leasing = request.POST.get('short_leasing')
		if short_leasing:
			houses = houses.filter(short_leasing=True)
		long_leasing = request.POST.get('long_leasing')
		if long_leasing:
			houses = houses.filter(long_leasing=True)
		houses = houses.filter(short_leasing_fee__gte=request.POST.get('short_min_fee'))
		houses = houses.filter(short_leasing_fee__lte=request.POST.get('short_max_fee'))
		houses = houses.filter(long_leasing_fee__gte=request.POST.get('long_min_fee'))
		houses = houses.filter(long_leasing_fee__lte=request.POST.get('long_max_fee'))
	return render(request, 'search_house.html', {'houses': houses, "username": login_name, 'kind': kind})


def specific_info(request):
	message = ""
	if request.method == 'POST':
		house = models.House.objects.get(house_id=request.POST.get('selected_id'))
		apps = models.Application.objects.all()
		apps = apps.filter(house_id=request.POST.get('selected_id')).filter(is_paid=True)
		if request.POST.get('type') == 'short':
			rent_type = 'short'
		elif request.POST.get('type') == 'long':
			rent_type = 'long'
		else:
			return render(request, 'specific_info.html', {'house': house, 'apps': apps})
		application = models.Application()
		user = models.User.objects.get(id=request.session.get('id'))
		application.username = user.username
		application.house_id = request.POST.get('selected_id')
		application.rent_type = rent_type
		application.s_y = request.POST.get('s_y')
		application.s_m = request.POST.get('s_m')
		application.s_d = request.POST.get('s_d')
		application.duration = request.POST.get('duration')
		application.save()
		if house.status == 2:
			house.status = 0
		house.save()
		message = '提交订单成功'
	return render(request, 'specific_info.html', {'house': house, 'apps': apps, 'message': message})


def add_technician(request):
	message = ""
	if request.method == 'POST':
		username = request.POST.get('username').strip()
		password1 = request.POST.get('password1').strip()
		password2 = request.POST.get('password2').strip()
		contact_number = request.POST.get('contact_number').strip()
		if username and password1 and password2:
			technician = models.Technician()
			if models.Technician.objects.filter(username=username).exists():
				message = '该用户名已被注册'
				return render(request, 'add_technician.html', {"message": message})
			if password1 != password2:
				message = '两次密码不一致'
				return render(request, 'add_technician.html', {"message": message})
			technician.username = username
			technician.password = password1
			technician.contact_number = contact_number
			technician.save()
			message = '工单师傅注册成功'
	return render(request, 'add_technician.html', {"message": message})


def check_app(request):
	apps = models.Application.objects.all()
	login_name = ""
	if request.session.get('id'):
		cs = models.Customerserver.objects.get(id=request.session.get('id'))
		login_name = cs.username
	if request.method == 'POST':
		if request.POST.get('not_allowed'):
			apps = apps.filter(is_allowed=False)
		if request.POST.get('apply_id'):
			apps = apps.filter(apply_id=request.POST.get('apply_id'))
		if request.POST.get('username'):
			apps = apps.filter(username=request.POST.get('username'))
		if request.POST.get('house_id'):
			apps = apps.filter(house_id=request.POST.get('house_id'))
	return render(request, 'check_app.html', {'apps': apps, "csname": login_name})


def confirm_payment(request):
	message = ""
	if request.method == 'POST':
		app = models.Application.objects.get(apply_id=request.POST.get('apply_id'))
		app.is_paid = True
		house = models.House.objects.get(house_id=app.house_id)
		house.status = 1
		app.save()
		house.save()
		message = "确认成功"
		return render(request, 'confirm_payment.html', {'message': message})


def user_info(request):
	login_name = ""
	message = ""
	if request.session.get('id'):
		cs = models.Customerserver.objects.get(id=request.session.get('id'))
		login_name = cs.username
	if request.POST.get('apply_id'):
		apply_id = request.POST.get('apply_id')
		request.session['apply_id'] = apply_id
	if request.method == 'POST':
		if request.POST.get('user_name'):
			user = models.User.objects.get(username=request.POST.get('user_name'))

		if request.POST.get('commit') != '1':
			return render(request, 'user_info.html', {'user': user, 'message':message})
		else:
			apply_id = request.session.get('apply_id')
			app = models.Application.objects.get(apply_id=apply_id)
			app.is_allowed = True
			app.allowed_by = login_name
			app.save()
			message = "完成审核"
			house = models.House.objects.get(house_id=app.house_id)
			user = models.User.objects.get(username=app.username)
			document = Document()
			p = document.add_heading('房屋租赁合同', 0)
			p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
			p = document.add_paragraph('甲方：青年租房有限公司')
			p = document.add_paragraph('乙方：' + str(user.real_name))
			p = document.add_paragraph('双方经友好协商，根绝《合同法》及国家、当地政府对房屋租赁的有关规定，就租赁房屋一事达成以下协议：')
			p.paragraph_format.first_line_indent = Inches(0.3)
			p = document.add_paragraph('')
			p.add_run('第一条 ').bold = True
			p.add_run('甲方保证向乙方出租的房屋为合法拥有完全的所有权和使用权，乙方确保了解该房屋全貌，并以现状为准。')
			p = document.add_paragraph('')
			p.add_run('第二条 ').bold = True
			p.add_run('出租房屋坐落于' + str(house.address) + '，室内附属设施齐全。')
			p = document.add_paragraph('')
			p.add_run('第三条 ').bold = True
			p.add_run(
				'租金每月人民币' + str(house.long_leasing_fee) + '元。' + '自' + str(app.s_y) + '年' + str(app.s_m) + '月' + str(
					app.s_d) + '日起开始入住，' + '共' + str(app.duration) + '个月，租金每月支付一次；自本合同生效之日起，'
				+ '乙方应先行向甲方支付一个月的租金，以后每月前一周内付清下月租金。租金以现金支付。')
			p = document.add_paragraph('')
			p.add_run('第四条 ').bold = True
			p.add_run('租赁期间，乙方因正常生活之需要的煤气费、水电费、暖气费、有线电视费、网络使用费等均由乙方承担。')
			p = document.add_paragraph('')
			p.add_run('第五条 ').bold = True
			p.add_run('租赁期间，乙方不得将房屋转租给第三方使用，否则，甲方有权单方面终止合同，收回房屋，且可以追究乙方的违约责任。')
			p = document.add_paragraph('')
			p.add_run('第六条 ').bold = True
			p.add_run('本协议一式两份，甲、乙两方各存一份，均具有同等法律效力。')
			p = document.add_paragraph('')
			p.add_run('第七条 ').bold = True
			p.add_run('本协议自双方签字之日起生效。')
			p = document.add_paragraph('甲方：       ')
			p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			p = document.add_paragraph('乙方：       ')
			p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			p = document.add_paragraph('签约时间：       年     月     日')
			p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			document.save('./static/applications/' + str(app.apply_id) + '.docx')
	return render(request, 'user_info.html', {'message':message})


def cancel_apply(request):
	if request.method == 'POST':
		if request.POST.get('apply_id'):
			app = models.Application.objects.get(apply_id=request.POST.get('apply_id'))
			app.delete()
			return render(request, 'check_app.html', {'message': '拒绝成功'})
	return render(request, 'cancel_apply.html', {'message': '删除错误'})


def delete_house(request):
	message = ""
	if request.method == 'POST':
		house = models.House.objects.get(house_id=request.POST.get('house_id'))
		if os.path.exists(settings.MEDIA_ROOT + "/demo1/" + str(house.house_id)):
			shutil.rmtree(settings.MEDIA_ROOT + "/demo1/" + str(house.house_id))
		house.delete()
		message = '下架成功'
	return render(request, 'search_house.html', {'message': message})


def my_app(request):
	my_apps = models.Application.objects.filter(username=models.User.objects.get(id=request.session.get('id')).username)
	login_name = ""
	if request.session.get('id'):
		user = models.User.objects.get(id=request.session.get('id'))
		login_name = user.username
	if request.method == 'POST':
		if request.POST.get('house_id'):
			my_apps = my_apps.filter(house_id=request.POST.get('house_id'))
	return render(request, 'my_app.html', {'my_apps': my_apps, "username": login_name})


def pay(request):
	if request.method == 'POST':
		house = models.House.objects.get(house_id=request.POST.get('house_id'))
		app = models.Application.objects.get(apply_id=request.POST.get('apply_id'))
		rent = house.short_leasing_fee * app.duration
	return render(request, 'pay.html', {'rent': rent})


def repairing(request):
	if request.method == 'POST':
		house_id = request.POST.get('house_id')
		if house_id:
			content = request.POST.get('content')
			files = request.FILES.getlist('picture', None)
			if files:
				for file in files:
					if not str(file.name).__contains__("jpg"):
						return render(request, 'repairing.html', {"message": "图片仅支持jpg格式", "house_id": house_id})
			else:
				return render(request, 'repairing.html', {"message": "请上传图片", "house_id": house_id})
			if not content:
				return render(request, 'repairing.html', {"message": "报修原因不能为空", "house_id": house_id})
			repairing = models.Repairing()
			repairing.house_id = house_id
			repairing.username = models.User.objects.get(id=request.session.get('id')).username
			repairing.content = content
			repairing.save()
			repairing.picture = "/demo1/" + str(house_id) + "_" + str(repairing.repair_id) + "/"
			os.makedirs(settings.MEDIA_ROOT + "/demo1/" + str(house_id) + "_" + str(repairing.repair_id) + "/")
			i = 1
			for file in files:
				filename = settings.MEDIA_ROOT + "/demo1/" + str(house_id) + "_" + str(repairing.repair_id) + "/" + str(i) + ".jpg"
				with open(filename, 'wb') as pic:
					for c in file.chunks():
						pic.write(c)
				i += 1
			repairing.pic_num = i - 1
			repairing.save()
			return render(request, 'my_app.html', {"message": "添加成功"})
		else:
			app = models.Application.objects.get(apply_id=request.POST.get('selected_id'))
			return render(request, 'repairing.html', {"house_id": app.house_id})
	return render(request, 'home.html', {"message": "错误"})


def manage_repairing(request):
	repairing = models.Repairing.objects.all()
	login_name = ""
	if request.session.get('id'):
		cs = models.Customerserver.objects.get(id=request.session.get('id'))
		login_name = cs.username
	if request.method == 'POST':
		if request.POST.get('not_allowed'):
			repairing = repairing.filter(is_allowed=False)
		if request.POST.get('repairing_id'):
			repairing = repairing.filter(repair_id=request.POST.get('repair_id'))
		if request.POST.get('username'):
			repairing = repairing.filter(username=request.POST.get('username'))
		if request.POST.get('house_id'):
			repairing = repairing.filter(house_id=request.POST.get('house_id'))
	return render(request, 'manage_repairing.html', {"login_name": login_name, "repairing": repairing})


def cancel_repairing(request):
	if request.method == 'POST':
		if request.POST.get('repairing_id'):
			repairing = models.Repairing.objects.filter(repair_id=request.POST.get('repairing_id')).first()
			new_message = models.MyMessage()
			new_message.receiver = repairing.username
			new_message.status = 1
			new_message.sender = models.Customerserver.objects.get(id=request.session.get('id')).username
			new_message.context = "您的报修已被管理员拒绝"
			new_message.title = "报修结果通知"
			new_message.save()
			repairing.delete()
			if os.path.exists(settings.MEDIA_ROOT + "/demo1/" + str(repairing.repair_id)):
				shutil.rmtree(settings.MEDIA_ROOT + "/demo1/" + str(repairing.repair_id))
			return render(request, 'main.html', {'message': '拒绝成功'})
	return render(request, 'manage_repairing.html', {'message': '拒绝错误'})


def arrange_repairing(request):
	if request.method == 'POST':
		if not request.POST.get('technician'):
			return render(request, 'manage_repairing.html', {'message': '请安排工单师傅'})
		if not models.Technician.objects.filter(username=request.POST.get('technician')).exists():
			return render(request, 'manage_repairing.html', {'message': '该工单师傅不存在'})
		repairing = models.Repairing.objects.filter(repair_id=request.POST.get('repairing_id')).first()
		repairing.repair_technician = models.Technician.objects.filter(username=request.POST.get('technician')).first().username
		repairing.is_allow = True
		repairing.save()
		new_message = models.MyMessage()
		new_message.receiver = repairing.username
		new_message.status = 1
		new_message.sender = models.Customerserver.objects.get(id=request.session.get('id')).username
		new_message.context = "您的报修已被管理员通过，请等待编号为" + str(repairing.repair_technician) + "的工单师傅进行处理。联系方式：" + str(models.Technician.objects.get(username=repairing.repair_technician).contact_number)
		new_message.title = "报修结果通知"
		new_message.save()
		return render(request, 'manage_repairing.html', {'message': '安排成功'})
	return render(request, 'manage_repairing.html', {'message': '安排错误'})


def my_repairing(request):
	if request.session.get('kind') == "user":
		my_repairing = models.Repairing.objects.filter(username=models.User.objects.get(id=request.session.get('id')).username)
		user = models.User.objects.get(id=request.session.get('id'))
		login_name = user.username
	else:
		my_repairing = models.Repairing.objects.filter(repair_technician=models.Technician.objects.get(id=request.session.get('id')).username)
		user = models.Technician.objects.get(id=request.session.get('id'))
		login_name = user.username
	if request.method == 'POST':
		if request.POST.get('house_id'):
			my_repairing = my_repairing.filter(house_id=request.POST.get('house_id'))
	return render(request, 'my_repairing.html', {'my_repairing': my_repairing, "username": login_name, "kind": request.session.get('kind')})


def comment(request):
	my_repairing = None
	if request.method == 'POST':
		repairing = models.Repairing.objects.get(repair_id=request.POST.get('selected_id'))
		repairing.comment = request.POST.get('comment')
		repairing.save()
		my_repairing = models.Repairing.objects.filter(username=models.User.objects.get(id=request.session.get('id')).username)
	return render(request, 'my_repairing.html', {'my_repairing': my_repairing, "message": "评价成功"})


def fix(request):
	my_repairing = None
	if request.method == 'POST':
		repairing = models.Repairing.objects.get(repair_id=request.POST.get('selected_id'))
		repairing.is_fix = True
		repairing.save()
		my_repairing = models.Repairing.objects.filter(username=models.User.objects.get(id=request.session.get('id')).username)
	return render(request, 'my_repairing.html', {'my_repairing': my_repairing, "message": "处理成功"})


def send_message(request):
	message = ""
	if request.method == 'POST':
		receiver = request.POST.get('username').strip()
		sender = request.session.get('id')
		title = request.POST.get('title').strip()
		context = request.POST.get('context')
		if not models.User.objects.filter(username=receiver).exists():
			message = "用户不存在"
			return render(request, 'send_message.html', {"message": message})
		elif not title or not context:
			message = "标题或内容不能为空"
			return render(request, 'send_message.html', {"message": message})
		else:
			message = "发送成功"
			new_message = models.MyMessage()
			new_message.receiver = receiver
			new_message.status = 1
			new_message.sender = models.Customerserver.objects.get(id=sender).username
			new_message.context = context
			new_message.title = title
			new_message.save()
	return render(request, 'send_message.html', {"message": message})


def my_message(request):
	message = ""
	user = models.User.objects.get(id=request.session.get('id'))
	my_messages = models.MyMessage.objects.filter(receiver=user)
	if request.method == 'POST':
		my_message_id = request.POST.get('selected_id')
		my_message = models.MyMessage.objects.filter(my_message_id=my_message_id).first()
		my_message.status = 0
		my_message.save()
		return render(request, 'specific_message.html', {"my_message": my_message})
	return render(request, 'my_message.html', {"message": message, "my_messages": my_messages})


def create_apply(request):
	message = ""
	cs = models.Customerserver.objects.get(id=request.session.get('id'))
	if request.method == 'POST':
		app = models.Application()
		app.username = request.POST.get('username')
		app.house_id = request.POST.get('house_id')
		app.rent_type = request.POST.get('rent_type')
		app.s_y = request.POST.get('s_y')
		app.s_m = request.POST.get('s_m')
		app.s_d = request.POST.get('s_d')
		app.duration = request.POST.get('duration')
		app.save()
		message = "创建成功"
	return render(request, 'create_apply.html', {'message': message})


# def export_app(request):
# 	login_name = ""
# 	if request.session.get('id'):
# 		user = models.User.objects.get(id=request.session.get('id'))
# 		login_name = user.username
# 	if request.method == 'POST':
# 		app = models.Application.objects.get(apply_id=request.POST.get('apply_id'))
# 		house = models.House.objects.get(house_id=request.POST.get('house_id'))
#
# 	return render(request, 'my_app.html', {'message': '导出成功'})


def download_app(request):
	app = models.Application.objects.get(apply_id=request.GET.get('apply_id'))
	file = open('./static/applications/'+str(app.apply_id)+'.docx', 'rb')
	response = FileResponse(file)
	response['Content-Type'] = 'application/msword'
	response['Content-Disposition'] = 'attachment;filename='+str(app.apply_id)+'.docx'
	return response


def remind(request):
	if request.method == 'POST':
		apps = models.Application.objects.all()
		login_name = ""
		if request.session.get('id'):
			cs = models.Customerserver.objects.get(id=request.session.get('id'))
			login_name = cs.username
		apps = apps.filter(is_paid=True)
		apps = apps.filter(rent_type='long')
		now = datetime.datetime.now().date()
		for app in apps:
			dt = datetime.date(app.s_y, app.s_m, app.s_d)
			delta = now - dt
			days = delta.days
			if days % 30 == 23:
				new_message = models.MyMessage()
				new_message.receiver = app.username
				new_message.status = 1
				new_message.sender = login_name
				house = models.House.objects.get(house_id=app.house_id)
				new_message.context = "您的订单ID为" + str(app.apply_id) + "的订单近期该缴费了。每月租金为：" + str(house.long_leasing_fee)
				new_message.title = "缴费通知"
				new_message.save()
				E.send_remind_email(models.User.objects.get(username=app.username).email, new_message.context)
		return render(request, 'check_app.html', {'apps': apps, "csname": login_name})


def check_user(request):
	users = models.User.objects.filter()
	if request.method == 'POST':
		if request.POST.get('delete_name'):
			user = models.User.objects.get(username=request.POST.get('delete_name'))
			user.delete()
			users = models.User.objects.filter()
			return render(request, 'check_user.html', {'users': users, 'message': "删除成功"})
		elif request.POST.get('selected_name'):
			user = models.User.objects.get(username=request.POST.get('selected_name'))
			return render(request, 'CS_show_info.html', {'user': user})
		else:
			if request.POST.get('username'):
				users = users.filter(username=request.POST.get('username'))
			if request.POST.get('user_id'):
				users = users.filter(id=request.POST.get('user_id'))
			return render(request, 'check_user.html', {'users': users})
	return render(request, 'check_user.html', {'users': users})


def CS_change_info(request):
	if request.method == 'POST':
		username = models.User.objects.get(id=request.session.get('id')).username
		password = models.User.objects.get(id=request.session.get('id')).password
		email = models.User.objects.get(id=request.session.get('id')).email
		real_name = models.User.objects.get(id=request.session.get('id')).real_name
		ID_card = models.User.objects.get(id=request.session.get('id')).ID_card
		models.User.objects.filter(username=username).delete()
		new_username = request.POST.get('username').strip()
		new_password1 = request.POST.get('password1').strip()
		new_password2 = request.POST.get('password2').strip()
		legal, message = checkout(new_username, new_password1, new_password2, email)
		if not legal:
			user = add_user(username, password, email, real_name, ID_card, False)
			return render(request, 'CS_show_info.html', locals())
		user = add_user(new_username, new_password1, email, real_name, ID_card, False)
		message = '修改成功'
		return render(request, 'check_user.html', locals())
	else:
		return render(request, 'CS_show_info.html', locals())